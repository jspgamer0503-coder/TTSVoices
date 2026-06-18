"""
file_extractor.py — TTS Voices 2.5.2

Maintained by the opencode AI assistant — see README.md.
Multi-format document text extraction: PDF, DOCX, DOC, EPUB, HTML, RTF,
ODT, ODT-encrypted, TXT, MD, CSV. Includes format detection, encoding
fallback chain, and password propagation for encrypted Office formats.
"""
import os, re, subprocess, tempfile, zipfile, base64, hashlib, zlib
from pathlib import Path
import bug_tracker


def _detect_encoding(raw: bytes, sample: int = 4096) -> str:
    """Fast encoding detection: BOM check, then UTF-8, then chardet."""
    # BOM checks (bytes literals safe for ASCII-only source)
    BOM_UTF8   = bytes([0xEF, 0xBB, 0xBF])
    BOM_UTF16L = bytes([0xFF, 0xFE])
    BOM_UTF16B = bytes([0xFE, 0xFF])
    if raw.startswith(BOM_UTF8):   return 'utf-8-sig'
    if raw.startswith(BOM_UTF16L): return 'utf-16-le'
    if raw.startswith(BOM_UTF16B): return 'utf-16-be'
    try:
        raw[:sample].decode('utf-8')
        return 'utf-8'
    except UnicodeDecodeError:
        pass
    try:
        import chardet
        return chardet.detect(raw[:sample]).get('encoding') or 'utf-8'
    except ImportError:
        return 'utf-8'


# Pre-compiled regex — ~40% faster XML extraction
_RE_PARA_VOID  = re.compile(r'<(text:p|text:h|text:list-item)[^>]*/>')
_RE_PARA_OPEN  = re.compile(r'<(text:p|text:h|text:list-item)[^>]*>')
_RE_PARA_CLOSE = re.compile(r'</(text:p|text:h|text:list-item)>')
_RE_TAB        = re.compile(r'<text:tab[^>]*/>')
_RE_LINEBREAK  = re.compile(r'<text:line-break[^>]*/>')
_RE_TAG        = re.compile(r'<[^>]+>')
# Whitespace run: SPACES only. Tabs are preserved as ODT <text:tab/> markers
# in the ODT pass and must not be squashed here.
_RE_WHITESPACE = re.compile(r' +')
_RE_NEWLINES   = re.compile(r'\n{3,}')

# ── Password dialog helper ────────────────────────────────────────────────────
def _ask_password(filename: str) -> str:
    """
    Thread-safe themed password dialog.
    Scheduled on the main Tk thread; calling thread blocks until user responds.
    The main textarea is disabled during the dialog to prevent stray keystrokes.
    """
    import threading as _threading
    import tkinter as tk

    root = getattr(tk, "_default_root", None)
    if root is None:
        bug_tracker.warning("No Tk root available for password dialog")
        return ""

    result = {"password": None}
    done   = _threading.Event()

    def _show_dialog():
        try:
            # ── Disable main textarea so keystrokes can't escape into it ──
            _textarea = None
            try:
                # Walk widget tree to find the Text widget
                def _find_text(w):
                    if isinstance(w, tk.Text):
                        return w
                    for child in w.winfo_children():
                        found = _find_text(child)
                        if found:
                            return found
                    return None
                _textarea = _find_text(root)
                if _textarea:
                    _textarea.configure(state="disabled")
            except Exception:
                pass

            win = tk.Toplevel(root)
            win.title("Password Required")
            win.configure(bg="#080c18")
            win.resizable(False, False)
            win.transient(root)
            win.attributes("-topmost", True)

            # Header
            hdr = tk.Frame(win, bg="#060a14")
            hdr.pack(fill="x")
            tk.Label(hdr, text="🔒  Password Required",
                     font=("Courier New", 12, "bold"),
                     fg="#00c8ff", bg="#060a14",
                     padx=20, pady=12).pack(side="left")
            tk.Frame(win, bg="#1a2a45", height=1).pack(fill="x")

            # Body
            body = tk.Frame(win, bg="#080c18", padx=24, pady=16)
            body.pack(fill="x")
            tk.Label(body, text="This file is password protected:",
                     font=("Courier New", 9),
                     fg="#8fa3c4", bg="#080c18").pack(anchor="w")
            tk.Label(body, text=Path(filename).name,
                     font=("Courier New", 10, "bold"),
                     fg="#dde4f0", bg="#080c18",
                     wraplength=340, justify="left").pack(anchor="w", pady=(2,12))
            tk.Label(body, text="Password:",
                     font=("Courier New", 9, "bold"),
                     fg="#8fa3c4", bg="#080c18").pack(anchor="w")

            pw_var   = tk.StringVar()
            pw_entry = tk.Entry(body, textvariable=pw_var,
                                show="●",
                                bg="#0d1526", fg="#dde4f0",
                                insertbackground="#00c8ff",
                                relief="flat",
                                font=("Courier New", 11),
                                width=32, bd=0,
                                highlightthickness=2,
                                highlightbackground="#1a6cf5",
                                highlightcolor="#00c8ff")
            pw_entry.pack(fill="x", ipady=8, pady=(4,4))

            show_var = tk.BooleanVar(value=False)
            def _toggle():
                pw_entry.configure(show="" if show_var.get() else "●")
            tk.Checkbutton(body, text="Show password",
                           variable=show_var, command=_toggle,
                           font=("Courier New", 8), fg="#64748b", bg="#080c18",
                           selectcolor="#0d1526", activebackground="#080c18",
                           activeforeground="#8fa3c4",
                           highlightthickness=0, bd=0,
                           cursor="hand2").pack(anchor="w", pady=(2,0))

            err_var = tk.StringVar(value="")
            tk.Label(body, textvariable=err_var,
                     font=("Courier New", 8, "bold"),
                     fg="#ef4444", bg="#080c18").pack(anchor="w", pady=(4,0))

            tk.Frame(win, bg="#1a2a45", height=1).pack(fill="x")

            foot = tk.Frame(win, bg="#0d1526", pady=12)
            foot.pack(fill="x")

            def _cleanup_and_signal(pw_value):
                """Re-enable textarea, release grab, signal done."""
                if _textarea:
                    try:
                        _textarea.configure(state="normal")
                    except Exception:
                        pass
                result["password"] = pw_value
                try:
                    win.grab_release()
                    win.destroy()
                except Exception:
                    pass
                done.set()

            def _ok(event=None):
                pw = pw_var.get().strip()
                if not pw:
                    err_var.set("⚠  Please enter a password")
                    pw_entry.focus_force()
                    return
                _cleanup_and_signal(pw)

            def _cancel(event=None):
                _cleanup_and_signal("")

            pw_entry.bind("<Return>", _ok)
            win.bind("<Escape>",      _cancel)
            win.protocol("WM_DELETE_WINDOW", _cancel)

            tk.Button(foot, text="  Cancel  ",
                      font=("Courier New", 9, "bold"),
                      bg="#111e33", fg="#8fa3c4", relief="flat",
                      padx=12, pady=6, cursor="hand2",
                      activebackground="#1a2a45",
                      command=_cancel).pack(side="right", padx=(4,16))
            tk.Button(foot, text="  Unlock  ",
                      font=("Courier New", 9, "bold"),
                      bg="#1a6cf5", fg="white", relief="flat",
                      padx=12, pady=6, cursor="hand2",
                      activebackground="#1d6aff",
                      command=_ok).pack(side="right", padx=4)

            # Geometry — centred
            win.update_idletasks()
            w  = win.winfo_reqwidth()
            h  = win.winfo_reqheight()
            sw = root.winfo_screenwidth()
            sh = root.winfo_screenheight()
            win.geometry(f"{w}x{h}+{(sw-w)//2}+{(sh-h)//2}")

            # Grab + focus — force immediately
            win.deiconify()
            win.lift()
            win.update()
            win.grab_set()
            pw_entry.focus_force()   # force_focus guarantees keyboard goes here

        except Exception as e:
            bug_tracker.error(f"Password dialog failed: {e}")
            result["password"] = ""
            done.set()

    # Use after_idle so the event loop is fully settled before showing dialog
    root.after_idle(_show_dialog)
    done.wait(timeout=120)
    return result["password"] or ""

# ── Plain text / Markdown ─────────────────────────────────────────────────────
def extract_txt(path: str) -> str:
    """Extract plain text using fast BOM-first encoding detection."""
    with open(path, "rb") as f:
        raw = f.read()
    enc = _detect_encoding(raw)
    try:
        return raw.decode(enc)
    except (UnicodeDecodeError, LookupError):
        return raw.decode("utf-8", errors="replace")

def extract_pdf(path: str, password: str = "") -> str:
    """Extract PDF text, prompting for password if encrypted."""

    # ── Try pikepdf first to detect/unlock encryption reliably ───────────
    try:
        import pikepdf
        try:
            pdf_obj = pikepdf.open(path, password=password) if password else pikepdf.open(path)
            pdf_obj.close()
        except pikepdf.PasswordError:
            pw = password or _ask_password(path)
            if not pw:
                raise RuntimeError("PDF is password-protected. Please provide the password.")
            try:
                pdf_obj = pikepdf.open(path, password=pw)
                pdf_obj.close()
            except pikepdf.PasswordError:
                raise RuntimeError("Incorrect PDF password.")
            # pikepdf verified `pw` — propagate it so pdfplumber/pypdf/pdftotext
            # don't trigger a second password prompt for the same file.
            password = pw
    except ImportError:
        pass  # pikepdf not installed, try other methods

    # ── pdfplumber ────────────────────────────────────────────────────────
    try:
        import pdfplumber
        kwargs = {"password": password} if password else {}
        try:
            with pdfplumber.open(path, **kwargs) as pdf:
                parts = [p.extract_text() for p in pdf.pages if p.extract_text()]
                if parts:
                    return "\n\n".join(parts)
        except Exception as e:
            err = str(e).lower()
            if "password" in err or "encrypt" in err or "decrypt" in err:
                pw = password or _ask_password(path)
                if not pw:
                    raise RuntimeError("PDF is password-protected.")
                # Retry inline; on failure, fall through to pypdf/pdftotext
                # instead of raising out of the whole extract_pdf call.
                try:
                    with pdfplumber.open(path, password=pw) as pdf:
                        parts = [p.extract_text() for p in pdf.pages if p.extract_text()]
                        if parts:
                            return "\n\n".join(parts)
                    password = pw
                except Exception as retry_err:
                    bug_tracker.warning(f"pdfplumber retry failed: {retry_err}")
            else:
                bug_tracker.warning(f"pdfplumber failed: {e}")
    except ImportError:
        pass

    # ── pypdf fallback ────────────────────────────────────────────────────
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        if reader.is_encrypted:
            pw = password or _ask_password(path)
            if not pw:
                raise RuntimeError("PDF is password-protected.")
            if reader.decrypt(pw) == 0:
                raise RuntimeError("Incorrect PDF password.")
            password = pw
        text = "\n\n".join(
            p.extract_text() for p in reader.pages if p.extract_text())
        if text.strip():
            return text
    except ImportError:
        pass

    # ── pdftotext CLI last resort ─────────────────────────────────────────
    try:
        args = ["pdftotext"]
        if password: args += ["-upw", password]
        args += [path, "-"]
        r = subprocess.run(args, capture_output=True, text=True, timeout=30)
        if r.returncode == 0 and r.stdout.strip():
            return r.stdout
    except Exception:
        pass

    raise RuntimeError("Cannot extract PDF. Install: pip install pdfplumber pikepdf")

# ── DOCX / Word ───────────────────────────────────────────────────────────────
def extract_docx(path: str, password: str = "") -> str:
    # Handle encrypted DOCX via msoffcrypto
    if password:
        try:
            import msoffcrypto, io
            with open(path, "rb") as f:
                office_file = msoffcrypto.OfficeFile(f)
                office_file.load_key(password=password)
                decrypted = io.BytesIO()
                office_file.decrypt(decrypted)
                decrypted.seek(0)
                from docx import Document
                doc = Document(decrypted)
                return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except msoffcrypto.exceptions.InvalidKeyError:
            # Wrong password — let the encrypted-detection branch below
            # re-prompt rather than silently falling through and possibly
            # raising a misleading "not a Word file" error later.
            raise
        except Exception as e:
            bug_tracker.warning(f"DOCX decrypt failed: {e}")

    # Check if encrypted without password provided
    try:
        import msoffcrypto
        with open(path, "rb") as f:
            of = msoffcrypto.OfficeFile(f)
            if of.is_encrypted():
                pw = _ask_password(path)
                if pw:
                    return extract_docx(path, pw)
                raise RuntimeError("DOCX is password-protected. Please provide the password.")
    except ImportError:
        pass
    except RuntimeError:
        raise

    # Normal DOCX
    try:
        from docx import Document
        try:
            doc = Document(path)
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        parts.append(row_text)
            return "\n".join(parts)
        except Exception as first:
            if "content type" in str(first).lower() or "not a Word file" in str(first).lower():
                # Template file - patch content type, then fall through to
                # the raw XML reader if the patched Document(tmp) also fails.
                tmp = tempfile.mktemp(suffix=".docx")
                try:
                    try:
                        with zipfile.ZipFile(path, "r") as zin, \
                             zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
                            for item in zin.infolist():
                                data = zin.read(item.filename)
                                if item.filename == "[Content_Types].xml":
                                    data = data.replace(
                                        b"wordprocessingml.template",
                                        b"wordprocessingml.document"
                                    )
                                zout.writestr(item, data)
                        doc = Document(tmp)
                        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                    except Exception:
                        # Patched template still unreadable — try raw XML
                        with zipfile.ZipFile(path, "r") as z:
                            xml = z.read("word/document.xml").decode("utf-8", errors="replace")
                        return re.sub(r'\s+', ' ', re.sub(r'<[^>]+>', ' ', xml)).strip()
                finally:
                    if os.path.exists(tmp): os.unlink(tmp)
            # Raw XML fallback
            try:
                with zipfile.ZipFile(path, "r") as z:
                    xml = z.read("word/document.xml").decode("utf-8", errors="replace")
                return re.sub(r'\s+', ' ', re.sub(r'<[^>]+>', ' ', xml)).strip()
            except Exception:
                raise RuntimeError(f"Cannot read Word file: {first}")
    except ImportError:
        raise RuntimeError("Install python-docx: pip install python-docx")

# ── DOC (legacy) ──────────────────────────────────────────────────────────────
def extract_doc(path: str) -> str:
    for cmd in [["antiword", path], ["catdoc", path]]:
        try:
            r = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            if r.returncode == 0 and r.stdout.strip():
                return r.stdout
            if r.returncode != 0:
                # Non-zero exit — log the tool's complaint so the "install
                # antiword" error at the end isn't misleading when the real
                # problem is an encrypted or corrupt .doc.
                tool = cmd[0]
                err_snip = (r.stderr or b"").decode("utf-8", errors="replace")[:200]
                bug_tracker.warning(f"{tool} rc={r.returncode}: {err_snip}")
        except FileNotFoundError:
            continue
        except subprocess.TimeoutExpired:
            # Try the next tool, or fall through to the final error
            continue
    raise RuntimeError("Cannot read .doc - install antiword: sudo apt install antiword")

# ── EPUB ──────────────────────────────────────────────────────────────────────
def extract_epub(path: str) -> str:
    try:
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup
        book  = epub.read_epub(path)
        parts = []
        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            soup = BeautifulSoup(item.get_content(), "html.parser")
            text = soup.get_text(separator="\n").strip()
            if text:
                parts.append(text)
        return "\n\n".join(parts)
    except ImportError:
        raise RuntimeError("Install: pip install ebooklib beautifulsoup4")

# ── HTML ──────────────────────────────────────────────────────────────────────
def extract_html(path: str) -> str:
    with open(path, "rb") as f:
        raw = f.read()
    html = raw.decode(_detect_encoding(raw), errors="replace")
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script","style","head","nav","footer","aside"]):
            tag.decompose()
        # Preserve paragraph breaks
        for tag in soup(["p","br","h1","h2","h3","h4","h5","h6","li","tr"]):
            tag.insert_before("\n")
        return re.sub(r'\n{3,}', '\n\n', soup.get_text()).strip()
    except ImportError:
        return re.sub(r'<[^>]+>', '', html)

# ── RTF ───────────────────────────────────────────────────────────────────────
def extract_rtf(path: str) -> str:
    try:
        from striprtf.striprtf import rtf_to_text
        with open(path, "r", errors="replace") as f:
            return rtf_to_text(f.read())
    except ImportError:
        raise RuntimeError("Install: pip install striprtf")

# ── ODT (full extraction with teletype) ──────────────────────────────────────

def extract_odt(path, password=""):
    try:
        with zipfile.ZipFile(path, "r") as z:
            mf_bytes = z.read("META-INF/manifest.xml")
    except Exception as e:
        raise RuntimeError(f"Cannot open ODT file: {e}")

    is_encrypted = b"encryption-data" in mf_bytes
    bug_tracker.info(f"ODT: encrypted={is_encrypted}")

    if not is_encrypted and not password:
        return _extract_odt_unencrypted(path)

    pw = password or _ask_password(path)
    if not pw:
        raise RuntimeError(
            "This ODT file is password-protected.\n"
            "Please provide the password to open it.")
    return _decrypt_odt_inline(path, pw, mf_bytes)




def _decrypt_odt_inline(path, password, mf_bytes):
    """
    Self-contained ODT decryption supporting both LibreOffice formats:

    LEGACY format (LO < 24.2 default): individual files encrypted (content.xml etc.)
      Manifest has: <file-entry full-path="content.xml"><encryption-data ...>
      Algorithm: SHA256(pw) → PBKDF2-HMAC-SHA1 → AES-256-CBC → zlib(-15)

    MODERN format (LO 24.2+ "wholesome encryption"): single outer package
      Manifest has: <file-entry full-path="encrypted-package"><encryption-data ...>
      Algorithm: SHA256(pw) → PBKDF2-SHA1 or Argon2id → AES-256-GCM
      Result: inner ZIP archive → read content.xml from it

    Based on odfdecrypt (Horsmann/odfdecrypt) reference implementation.
    """
    import xml.etree.ElementTree as ET

    MNS = "urn:oasis:names:tc:opendocument:xmlns:manifest:1.0"
    LNS = "urn:org:documentfoundation:names:experimental:office:xmlns:loext:1.0"
    A   = "{" + MNS + "}"
    L   = "{" + LNS + "}"

    # ── Parse manifest ─────────────────────────────────────────────────────
    params       = None
    format_type  = "legacy"   # "legacy" or "modern"
    target_file  = None

    try:
        root = ET.fromstring(mf_bytes)
        for entry in root.findall(".//{*}file-entry"):
            fp = (entry.get(A + "full-path") or
                  entry.get("full-path") or "")
            # Normalise path
            fp_norm = fp.lstrip("./").lstrip("/") or fp

            # Check for modern format (encrypted-package)
            if fp_norm == "encrypted-package" or fp == "encrypted-package":
                enc_el = entry.find(A + "encryption-data")
                if enc_el is not None:
                    format_type = "modern"
                    target_file = "encrypted-package"
                    params = _parse_enc_entry(entry, enc_el, A, L, is_modern=True)
                    break

            # Legacy format (content.xml)
            if fp_norm == "content.xml":
                enc_el = entry.find(A + "encryption-data")
                if enc_el is not None:
                    target_file = fp  # keep original path for zip read
                    params = _parse_enc_entry(entry, enc_el, A, L, is_modern=False)
                    # Don't break - modern format takes priority if found later

    except Exception as e:
        bug_tracker.warning(f"ODT ET parse: {e}")

    # Regex fallback for legacy format only
    if params is None:
        mf_str = mf_bytes.decode("utf-8", errors="replace")
        # Check if it's modern format (encrypted-package present)
        if "encrypted-package" in mf_str:
            format_type = "modern"
            target_file = "encrypted-package"
        else:
            format_type = "legacy"
            target_file = "content.xml"

        B = r"[A-Za-z0-9+/=]+"
        def rx(pat):
            m = re.search(pat, mf_str)
            return m.group(1) if m else None

        iv_str   = rx(r'initialisation-vector=["\'](' + B + r')["\']')
        salt_str = rx(r'(?:[\s"\'<])salt=["\'](' + B + r')["\']')
        itr_str  = rx(r'iteration-count=["\'](\d+)["\']')
        ksz_str  = rx(r'key-size=["\'](\d+)["\']')
        ck_str   = rx(r'(?:[\s"\'<])checksum=["\'](' + B + r')["\']')
        algo_str = rx(r'algorithm-name=["\']([^"\']+)["\']')
        kdf_str  = rx(r'key-derivation-name=["\']([^"\']+)["\']')
        sk_str   = rx(r'start-key-generation-name=["\']([^"\']+)["\']')

        if iv_str and salt_str:
            params = {
                "iv":         base64.b64decode(iv_str),
                "salt":       base64.b64decode(salt_str),
                "checksum":   base64.b64decode(ck_str) if ck_str else b"",
                "iterations": int(itr_str) if itr_str else 100000,
                "key_size":   int(ksz_str) if ksz_str else 32,
                "algo":       algo_str or "aes256-cbc",
                "kdf":        kdf_str or "PBKDF2",
                "start_key":  sk_str or "sha256",
                "argon2_t":   3, "argon2_m": 65536, "argon2_p": 4,
            }

    if not params or not params.get("iv") or not params.get("salt"):
        raise RuntimeError(
            "Could not read encryption parameters from this ODT.\n"
            "Run decrypt_test.py on your file to diagnose:\n"
            "  python3 decrypt_test.py yourfile.odt yourpassword")

    bug_tracker.info(
        f"ODT format={format_type} iv={params['iv'].hex()[:10]} "
        f"salt={params['salt'].hex()[:10]} iters={params.get('iterations','?')} "
        f"algo={params.get('algo','?')}")

    # ── Read encrypted bytes ───────────────────────────────────────────────
    enc_data = None
    with zipfile.ZipFile(path, "r") as z:
        for cp in (target_file,
                   "content.xml", "./content.xml",
                   "encrypted-package"):
            try:
                enc_data = z.read(cp)
                break
            except KeyError:
                continue
    if enc_data is None:
        raise RuntimeError("Encrypted data not found in archive — file may be corrupt.")

    # ── Crypto helpers ─────────────────────────────────────────────────────
    def aes_cbc(key, iv, ct):
        try:
            from Crypto.Cipher import AES
            return AES.new(key, AES.MODE_CBC, iv).decrypt(ct)
        except ImportError:
            from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
            from cryptography.hazmat.backends import default_backend
            c = Cipher(algorithms.AES(key), modes.CBC(iv), backend=default_backend())
            d = c.decryptor()
            return d.update(ct) + d.finalize()

    def aes_gcm(key, iv, ct):
        # ODF 1.3 "encrypted-package" format does NOT prepend the IV to the
        # ciphertext — the IV is supplied separately via the manifest. The
        # old "ct[:len(iv)] == iv" check was a leftover heuristic that
        # silently corrupted ciphertext on accidental byte matches.
        try:
            from Crypto.Cipher import AES
            tag, payload = ct[-16:], ct[:-16]
            return AES.new(key, AES.MODE_GCM, nonce=iv).decrypt_and_verify(payload, tag)
        except ImportError:
            from cryptography.hazmat.primitives.ciphers.aead import AESGCM
            return AESGCM(key).decrypt(iv, ct, None)

    def pbkdf2_sha1(sk, salt, n, ksz):
        try:
            from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
            from cryptography.hazmat.primitives import hashes
            from cryptography.hazmat.backends import default_backend
            k = PBKDF2HMAC(algorithm=hashes.SHA1(), length=ksz,
                            salt=salt, iterations=n, backend=default_backend())
            return k.derive(sk)
        except ImportError:
            return hashlib.pbkdf2_hmac("sha1", sk, salt, n, dklen=ksz)

    def argon2_key(start_key, params):
        from argon2.low_level import hash_secret_raw, Type
        return hash_secret_raw(
            secret=start_key, salt=params["salt"],
            time_cost=params.get("argon2_t", 3),
            memory_cost=params.get("argon2_m", 65536),
            parallelism=params.get("argon2_p", 4),
            hash_len=params.get("key_size", 32),
            type=Type.ID)

    # ── Generate start key from password ──────────────────────────────────
    pw_b     = password.encode("utf-8")
    sha256k  = hashlib.sha256(pw_b).digest()
    sha1k    = hashlib.sha1(pw_b).digest()
    kdf_name = params.get("kdf", "PBKDF2").lower()
    algo     = params.get("algo", "aes256-cbc").lower()
    expected = params.get("checksum", b"")
    salt     = params["salt"]
    n        = params.get("iterations", 100000)
    ksz      = params.get("key_size", 32)
    iv       = params["iv"]

    # ── MODERN FORMAT: Argon2id / AES-GCM ────────────────────────────────
    if "argon2" in kdf_name:
        try:
            key = argon2_key(sha256k, params)
            plain = aes_gcm(key, iv, enc_data)
            # Decompress the inner ODF archive
            try:
                inner_zip = zlib.decompress(plain, -15)
            except Exception:
                inner_zip = plain
            # Open inner ZIP and extract content.xml
            import io as _io
            with zipfile.ZipFile(_io.BytesIO(inner_zip)) as inner_z:
                xml_bytes = inner_z.read("content.xml")
            text = _xml_to_text(xml_bytes.decode("utf-8", errors="replace"))
            if text.strip():
                bug_tracker.info(f"ODT modern/Argon2 decrypted: {len(text):,} chars")
                return text
        except Exception as e:
            raise RuntimeError(
                f"Incorrect password (Argon2/GCM: {e})\n"
                "Please check your password and try again.")

    # ── MODERN FORMAT: PBKDF2 / AES-GCM (non-Argon2) ────────────────────
    if "gcm" in algo:
        for sk in (sha256k, sha1k, pw_b):
            try:
                key   = pbkdf2_sha1(sk, salt, n, ksz)
                plain = aes_gcm(key, iv, enc_data)
                try:
                    inner_zip = zlib.decompress(plain, -15)
                except Exception:
                    inner_zip = plain
                import io as _io
                with zipfile.ZipFile(_io.BytesIO(inner_zip)) as inner_z:
                    xml_bytes = inner_z.read("content.xml")
                text = _xml_to_text(xml_bytes.decode("utf-8", errors="replace"))
                if text.strip():
                    bug_tracker.info(f"ODT modern/GCM decrypted: {len(text):,} chars")
                    return text
            except Exception:
                continue

    # ── LEGACY FORMAT: all PBKDF2 / AES-CBC strategies ───────────────────
    strategies = [
        (sha256k, "sha1"),   # LO 24.x-26.x CONFIRMED
        (sha256k, "sha256"),
        (sha1k,   "sha1"),
        (sha1k,   "sha256"),
        (pw_b,    "sha256"),
        (pw_b,    "sha1"),
    ]

    for sk, prf in strategies:
        try:
            key = pbkdf2_sha1(sk, salt, n, ksz) if prf == "sha1" else \
                  hashlib.pbkdf2_hmac("sha256", sk, salt, n, dklen=ksz)
            dec = aes_cbc(key, iv, enc_data)
            pad = dec[-1] if dec else 0
            # Validate PKCS#7 padding properly. Wrong passwords often yield
            # a final byte in 1..16 by chance, which the old check passed
            # through and then zlib failed on. Comparing the last `pad` bytes
            # against the expected repetition is the standard PKCS#7 test.
            if not (1 <= pad <= 16) or dec[-pad:] != bytes([pad]) * pad:
                continue
            unc = dec[:-pad]

            ck_ok = False
            if expected:
                ck_ok = (hashlib.sha256(dec[:1024]).digest() == expected or
                         hashlib.sha256(unc[:1024]).digest() == expected)

            xml_bytes = None
            for wbits in (-15, 15):
                try:
                    cand = zlib.decompress(unc, wbits)
                    if cand.lstrip(b"\xef\xbb\xbf\x00")[:5] in (
                            b"<?xml", b"<offi", b"<text", b"<mani"):
                        xml_bytes = cand
                        break
                except Exception:
                    pass

            if ck_ok or xml_bytes is not None:
                if xml_bytes is None:
                    xml_bytes = zlib.decompress(unc, -15)

                # If result is a ZIP (modern inner archive), unwrap it
                if xml_bytes[:2] == b"PK":
                    import io as _io
                    try:
                        with zipfile.ZipFile(_io.BytesIO(xml_bytes)) as inner_z:
                            xml_bytes = inner_z.read("content.xml")
                    except Exception:
                        pass

                text = _xml_to_text(xml_bytes.decode("utf-8", errors="replace"))
                if text.strip():
                    bug_tracker.info(f"ODT legacy/{prf} decrypted: {len(text):,} chars")
                    return text
        except Exception:
            continue

    raise RuntimeError(
        "Incorrect password — the file could not be decrypted.\n"
        "Please check your password and try again.")


def _parse_enc_entry(entry, enc_el, A, L, is_modern):
    """Parse encryption parameters from a manifest file-entry element."""
    algo_el = enc_el.find(A + "algorithm")
    kd_el   = enc_el.find(A + "key-derivation")
    sk_el   = enc_el.find(A + "start-key-generation")

    def ga(el, attr):
        if el is None: return ""
        return el.get(A + attr, el.get(attr, ""))

    iv_b64   = ga(algo_el, "initialisation-vector")
    salt_b64 = ga(kd_el,   "salt")
    ck_b64   = ga(enc_el,  "checksum")
    kdf_name = ga(kd_el,   "key-derivation-name")
    algo_name= ga(algo_el, "algorithm-name")
    iters    = int(ga(kd_el, "iteration-count") or "100000")
    key_sz   = int(ga(kd_el, "key-size") or ("32" if is_modern else "16"))
    sk_algo  = ga(sk_el, "start-key-generation-name")

    if not iv_b64 or not salt_b64:
        return None

    return {
        "iv":        base64.b64decode(iv_b64),
        "salt":      base64.b64decode(salt_b64),
        "checksum":  base64.b64decode(ck_b64) if ck_b64 else b"",
        "iterations": iters,
        "key_size":   key_sz,
        "algo":       algo_name,
        "kdf":        kdf_name,
        "start_key":  sk_algo,
        "argon2_t":   int(kd_el.get(L + "argon2-iterations", "3"))    if kd_el is not None else 3,
        "argon2_m":   int(kd_el.get(L + "argon2-memory",     "65536")) if kd_el is not None else 65536,
        "argon2_p":   int(kd_el.get(L + "argon2-lanes",      "4"))    if kd_el is not None else 4,
    }



def _extract_odt_unencrypted(path):
    skip = {"styles.xml", "meta.xml", "settings.xml", "manifest.rdf", "mimetype"}
    candidates = ["content.xml", "./content.xml"]
    try:
        with zipfile.ZipFile(path, "r") as z:
            names = z.namelist()
            bug_tracker.info(f"ODT zip: {names[:8]}")
            for n in names:
                if n.endswith(".xml") and n not in skip and n not in candidates:
                    candidates.append(n)
            for cand in candidates:
                if cand not in names:
                    continue
                try:
                    data = z.read(cand)
                    if data[:2] == b"PK":
                        continue
                    text = _xml_to_text(data.decode("utf-8", errors="replace"))
                    if text.strip():
                        bug_tracker.info(f"ODT unencrypted: {len(text):,} chars")
                        return text
                except Exception:
                    continue
    except Exception as e:
        bug_tracker.warning(f"ODT unencrypted failed: {e}")
    raise RuntimeError("Could not extract text from ODT — file may be empty or corrupt.")




def _xml_to_text(xml: str) -> str:
    """Convert ODT XML to plain text using pre-compiled regex patterns."""
    xml = _RE_PARA_VOID.sub('\n', xml)
    xml = _RE_PARA_OPEN.sub('\n', xml)
    xml = _RE_PARA_CLOSE.sub('\n', xml)
    xml = _RE_TAB.sub('\t', xml)
    xml = _RE_LINEBREAK.sub('\n', xml)
    xml = _RE_TAG.sub('', xml)
    xml = (xml.replace('&amp;', '&').replace('&lt;', '<')
               .replace('&gt;', '>').replace('&quot;', '"')
               .replace('&apos;', "'").replace('&#160;', ' '))
    xml = _RE_WHITESPACE.sub(' ', xml)
    xml = _RE_NEWLINES.sub('\n\n', xml)
    return xml.strip()


def extract_csv(path: str) -> str:
    import csv
    rows = []
    try:
        with open(path, newline="", encoding="utf-8", errors="replace") as f:
            for row in csv.reader(f):
                rows.append(", ".join(cell for cell in row if cell.strip()))
    except Exception as e:
        bug_tracker.warning(f"CSV read error: {e}")
    return "\n".join(rows)

# ── Image / OCR ────────────────────────────────────────────────────────────────

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tiff", ".tif", ".webp"}

def _ensure_tesseract():
    """Ensure tesseract binary is available. Returns path to tesseract binary."""
    import shutil

    # 1. Check system PATH first
    sys_tess = shutil.which("tesseract")
    if sys_tess:
        return sys_tess

    # 2. Check local download
    local_tess = Path.home() / ".ttsvoices" / "tesseract" / "extracted" / "usr" / "bin" / "tesseract"
    if local_tess.exists():
        # Add its lib dir to LD_LIBRARY_PATH
        lib_dir = local_tess.parent.parent / "lib" / "x86_64-linux-gnu"
        if lib_dir.exists():
            os.environ.setdefault("LD_LIBRARY_PATH", "")
            if str(lib_dir) not in os.environ["LD_LIBRARY_PATH"]:
                os.environ["LD_LIBRARY_PATH"] = str(lib_dir) + ":" + os.environ["LD_LIBRARY_PATH"]
        # Set TESSDATA_PREFIX
        tessdata = local_tess.parent.parent / "share" / "tesseract-ocr" / "5" / "tessdata"
        if tessdata.exists():
            os.environ["TESSDATA_PREFIX"] = str(tessdata)
        os.environ["PATH"] = str(local_tess.parent) + ":" + os.environ.get("PATH", "")
        return str(local_tess)

    # 3. Try auto-download
    try:
        from dep_installer import _download_tesseract
        ok, path = _download_tesseract()
        if ok:
            return path
    except Exception:
        pass

    return None


def extract_image(path: str) -> str:
    """Extract text from an image using Tesseract OCR (pytesseract)."""
    # Check pytesseract is importable
    try:
        import pytesseract
    except ImportError:
        raise RuntimeError(
            "OCR is not installed.\n\n"
            "Install it with:\n"
            "  pip install pytesseract Pillow --break-system-packages\n\n"
            "Or open Settings → Check for package updates → Install All."
        )

    # Ensure tesseract binary is available
    tess_path = _ensure_tesseract()
    if not tess_path:
        raise RuntimeError(
            "Tesseract OCR engine is not installed.\n\n"
            "Open Settings → Updates → Install All to auto-install it, or run:\n"
            "  sudo apt install tesseract-ocr  (Linux)"
        )

    from PIL import Image
    bug_tracker.info(f"Running OCR on: {path}")
    try:
        img = Image.open(path)
        text = pytesseract.image_to_string(img)
        return text
    except Exception as e:
        raise RuntimeError(f"OCR failed on image: {e}")

# ── Registry ──────────────────────────────────────────────────────────────────
EXTRACTORS = {
    ".pdf":  extract_pdf,
    ".docx": extract_docx,
    ".doc":  extract_doc,
    ".epub": extract_epub,
    ".html": extract_html,
    ".htm":  extract_html,
    ".rtf":  extract_rtf,
    ".odt":  extract_odt,
    ".txt":  extract_txt,
    ".md":   extract_txt,
    ".csv":  extract_csv,
}
# Add image formats
for _ext in IMAGE_EXTS:
    EXTRACTORS[_ext] = extract_image

SUPPORTED_EXTENSIONS = list(EXTRACTORS.keys())
SUPPORTED_DISPLAY    = "PDF, DOCX, DOC, EPUB, HTML, RTF, ODT, TXT, MD, CSV, PNG, JPG, BMP, GIF, TIFF, WEBP"

def extract_text(path: str) -> str:
    """Extract text from any supported file. Handles passwords and locked files."""
    ext  = Path(path).suffix.lower()
    func = EXTRACTORS.get(ext)
    if func is None:
        raise ValueError(f"Unsupported format: {ext}. Supported: {SUPPORTED_DISPLAY}")

    # Check if file is readable (not locked by another app)
    try:
        with open(path, "rb") as f:
            f.read(16)  # try reading a few bytes
    except PermissionError:
        raise RuntimeError(
            f"Cannot open '{Path(path).name}' — the file may be open in another application.\n"
            "Please close it in LibreOffice / your document reader and try again."
        )
    except OSError as e:
        raise RuntimeError(f"Cannot access file: {e}")

    bug_tracker.info(f"Extracting text from {ext} file: {path}")
    return func(path)
